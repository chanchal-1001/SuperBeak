import tkinter as tk
from tkinter import filedialog
import fitz
import docx
import openai
import os
from langchain.indexes import VectorstoreIndexCreator #uses OpenAIEmbeddings

#from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
#from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
from tkinter import scrolledtext
import random
key = '<give your open ai key here>'
os.environ["OPENAI_API_KEY"] = key
openai.api_key = os.getenv("OPENAI_API_KEY")

class FileOperations:
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("Doc Reader")
        self.window.geometry("500x500")
        self.window.geometry("-0-0")
        self.seletct_button = tk.Button( text="Select Directory",font = "consolas 14" ,command=self.select_directory)
        self.seletct_button.pack() 
        # Create a label to display the selected file paths
        self.selected_dir = tk.Label( text="No directory selected!")
        self.selected_dir.pack()  
        self.selected_files = tk.Label( text="")
        self.selected_files.pack()
        self.selected_files_list = []
        self.upload_Button = tk.Button( text="Upload")
        self.hideButton(self.upload_Button)
        
        self.chat_Button = tk.Button( text="Start quering on your docs", )
        self.chat_Button.config(command=self.getVectorIndexForLoader)        
        # Create a scrolled text widget for displaying messages with a vertical scrollbar                
        #self.chat_box.pack()  
        self.errorMsg = tk.Label(text = "")
        self.errorMsg.pack()              
        
        self.window.mainloop()

    def select_directory(self):
        self.directory_path = filedialog.askdirectory(title="Select files")        
        if self.directory_path:  
            print(self.directory_path)
            self.selected_dir.config(text=f"Selected directory: {self.directory_path}")        
            self.files = os.listdir(self.directory_path)

            filesList = ''
            allowed_file_type = ['txt','pdf','docx']            
            for index, value in enumerate(self.files):
                if value.split('.')[1] in allowed_file_type:
                    filesList += f'\n {index+1}: {value}'
                    self.selected_files_list.append(value)

            #self.selected_files.config( text=newline.join([f'{index+1}: {value}' for index, value in enumerate(self.files)]))  
            self.selected_files.config( text= filesList )  
            self.showButton(self.upload_Button)
            self.upload_Button.config(command=self.upload_files)
        #self.upload_Button.pack() 
        
    def hideButton(self, button):
        button.pack_forget()
    
    def showButton(self, button):
        button.pack()
    
    def close_chatwindow(self):
        self.window.destroy()
    
    def showChatWindow(self):
        # Start the Tkinter event loop
        self.hideButton(self.chat_Button)
        
        # Create a scrolled text widget for displaying messages with a vertical scrollbar
        self.chat_box = scrolledtext.ScrolledText(self.window, wrap=tk.WORD, width=40, height=15)
        #chat_box.grid(row=0, column=0, padx=10, pady=10, columnspan=2)

        # Configure tags for different message styles (user and chatbot)
        self.chat_box.tag_configure("user", foreground="blue")
        self.chat_box.tag_configure("chatbot", foreground="green")
        self.chat_box.pack()
        # Create an Entry widget for user input
        self.entry = tk.Entry(self.window, width=40)
        #self.entry.grid(row=1, column=0, padx=10, pady=(0, 10))
        # Bind the Enter key to the send_message function
        self.entry.bind("<Return>", self.send_message)
        self.entry.pack()
        # Create a button to send messages
        self.send_button = tk.Button(self.window, text="Send", command=self.on_send_button_click)
        #self.send_button.grid(row=1, column=1, padx=10, pady=(0, 10))
        self.send_button.pack()
       
       
        # Configure row and column weights to make the text box resizable
        self.window.columnconfigure(0, weight=1)
        self.window.rowconfigure(0, weight=1)

        #self.window.mainloop()

    
    def upload_files(self):    
        corpus = ''
        count = 0
        allowed_file_type = ['txt','pdf','docx']
        for file_name in self.selected_files_list:                        
            if file_name.split('.')[1] in allowed_file_type:
                f_path = os.path.join(self.directory_path, file_name)
                print("Reading file : ", f_path)
                prefix = ' Content under ' + file_name + ' file : '
                if os.path.isfile(f_path)  and os.path.exists(f_path):                    
                    if f_path.endswith('.txt'):                                
                        file_content = self.read_txt(f_path)
                        corpus += prefix + file_content
                        count  = count + 1 
                    elif f_path.endswith('.docx'):   
                        file_content = self.read_docx(f_path)
                        corpus += prefix + file_content
                        count +=1
                    elif f_path.endswith('.pdf'):   
                        file_content = self.read_pdf(f_path)
                        corpus += prefix + file_content
                        count +=1
            else:
                print("Supported file types : docx, txt, pdf")
        self.corpus = corpus + " Total number of files are " + str(count) +"."
        print(self.corpus)
        self.hideButton(self.upload_Button)
        self.showButton(self.chat_Button)

    def read_docx(self,f_path):
        file_content = ""
        try:
            file_properties = self.get_file_properties(f_path)                
            # Create a Document object from the .docx file
            doc = docx.Document(f_path)
            # Iterate through paragraphs and append text to the content string
            for paragraph in doc.paragraphs:                               
                if paragraph.text:                   
                    file_content += paragraph.text + "\n"                
        except Exception as e:
            print(f"An error occurred: {e} reading the file : {f_path}")
            self.errorMsg.cofig(text = f'Problem occured reading the file : {f_path}')
        if int(file_properties['file_size']) > 0 and file_content == "":
            self.errorMsg.cofig(text = f'Problem occured reading the file : {f_path}')
        return file_content
    
    
    def read_txt(self, f_path):        
        file_content = ""
        try:
            file_properties = self.get_file_properties(f_path)                 
            with open(f_path, "r") as file:
                for line in file:
                    file_content += line 
        except Exception as e:
            print(f"An error occurred: {e} reading the file : {f_path}")
            self.errorMsg.cofig(text = f'Problem occured reading the file : {f_path}')
        if int(file_properties['file_size']) > 0 and file_content == "":
            self.errorMsg.cofig(text = f'Problem occured reading the file : {f_path}')
        return file_content
    
    
    def read_pdf(self, f_path):
        file_content = ''
        try:
            file_properties = self.get_file_properties(f_path) 
            # Open the PDF file
            pdf_document = fitz.open(f_path)
            # Iterate through pages and extract text

            for page_number in range(pdf_document.page_count):
                page = pdf_document[page_number]
                file_content += page.get_text()

            # Close the PDF file
            pdf_document.close()
        except Exception as e:
            print(f"An error occurred: {e} reading the file : {f_path}")
            self.errorMsg.cofig(text = f'Problem occured reading the file : {f_path}')
        if int(file_properties['file_size']) > 0 and file_content == "":
            self.errorMsg.cofig(text = f'Problem occured reading the file : {f_path}')
        return file_content
    
    def get_file_properties(self,file_path):
        # Get the file size in bytes
        file_size = os.path.getsize(file_path)

        # Get other file properties
        file_properties = {
            "file_size": file_size,            
            "isDirectory": os.path.isdir(file_path),
            "isFile": os.path.isfile(file_path),
            "lastModifiedTime": os.path.getmtime(file_path),
            "creationTime": os.path.getctime(file_path),
        }
        for key, value in file_properties.items():
                print(f"{key}: {value}")
        return file_properties
    
    def process_query(self):
        user_message = self.entry.get()
        print("User query:", user_message)
        if user_message:
            self.chat_box.insert(tk.END, f"You: {user_message}\n", "user")
            self.entry.delete(0, tk.END)  # Clear the entry widget
            # Simulate a chatbot response (replace this with your actual chatbot logic)
            chatbot_response = self.generate_chatbot_response(user_message)
            self.chat_box.insert(tk.END, f"Chatbot: {chatbot_response}\n", "chatbot")
            self.chat_box.yview(tk.END)  # Auto-scroll to the latest message
            
    def send_message(self, event):
        self.process_query()            
    
    def on_send_button_click(self):
        self.process_query()   
        
    def generate_chatbot_response(self,user_message):
        # Sample chatbot response logic (replace with your own logic)       
        if user_message:
            response = self.vector_index.query(user_message) 
            #response = ['Hi! how are you', "Tell me more.", "That's interesting!"]
            #response = random.choice(responses)
        return response

        
    def get_corpus_chunks(self,corpus):
        try:
            
            docs = []
            print("get_corpus_chunks calleds for text len : ",len(corpus))
            text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=5)
            texts = text_splitter.split_text(self.corpus)         
            docs = [Document(page_content = t) for t in texts]
            print("docs lenth : ", len(docs))    
            return docs
        except Exception as e:
            print(f"An error occurred during chunking: {e}")
            
            
    def getVectorIndexForLoader(self):
        try:
            corpus = self.corpus
            self.vector_index = VectorstoreIndexCreator().from_documents(self.get_corpus_chunks(corpus)) 
            self.showChatWindow()
        except Exception as e:
            print(f"An error occurred: {e}")
            self.errorMsg.config( text="Application is facing difficulty in connecting with Open.Ai."\
                                       " \n Consider checking your plan and billing details OR Try again later.")
            self.hideButton(self.chat_Button)
                        
if __name__ == "__main__":
    FileOperations()   
